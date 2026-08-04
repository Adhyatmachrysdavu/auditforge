"""Render laporan ke DOCX (D15) dengan letterhead brand.

Deterministik: menerima `ReportData` → mengembalikan byte .docx. Judul bagian
dilokalkan (ID/EN); isi naratif memakai bahasa saat draf dibuat.
"""
from __future__ import annotations

import io

from docx import Document
from docx.document import Document as DocumentT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.reporting.report_data import ReportData

_L = {
    "id": {
        "prepared_by": "Disiapkan oleh",
        "client": "Klien",
        "generated": "Dibuat",
        "posture": "Postur Keamanan",
        "exec_summary": "Ringkasan Eksekutif",
        "overview": "Gambaran Umum",
        "key_risks": "Risiko Utama",
        "recommendations": "Rekomendasi",
        "findings": "Temuan",
        "total": "Total temuan disetujui",
        "by_severity": "Per keparahan",
        "severity": "Keparahan",
        "priority": "Prioritas",
        "status": "Status",
        "description": "Deskripsi",
        "impact": "Dampak",
        "recommendation": "Rekomendasi",
        "edited": "disunting auditor",
        "empty": "Tidak ada temuan disetujui untuk dilaporkan.",
        "confidential": "RAHASIA — hanya untuk penerima yang berwenang.",
    },
    "en": {
        "prepared_by": "Prepared by",
        "client": "Client",
        "generated": "Generated",
        "posture": "Security Posture",
        "exec_summary": "Executive Summary",
        "overview": "Overview",
        "key_risks": "Key Risks",
        "recommendations": "Recommendations",
        "findings": "Findings",
        "total": "Total approved findings",
        "by_severity": "By severity",
        "severity": "Severity",
        "priority": "Priority",
        "status": "Status",
        "description": "Description",
        "impact": "Impact",
        "recommendation": "Recommendation",
        "edited": "edited by auditor",
        "empty": "No approved findings to report.",
        "confidential": "CONFIDENTIAL — for authorized recipients only.",
    },
}

_POSTURE = {
    "id": {
        "critical": "Kritis", "elevated": "Tinggi", "moderate": "Sedang",
        "low": "Rendah", "clean": "Bersih",
    },
    "en": {
        "critical": "Critical", "elevated": "Elevated", "moderate": "Moderate",
        "low": "Low", "clean": "Clean",
    },
}


def _accent(color_hex: str) -> RGBColor:
    try:
        return RGBColor.from_string(color_hex.lstrip("#"))
    except (ValueError, AttributeError):
        return RGBColor.from_string("1E5F9F")


def render_docx(data: ReportData, *, accent: str = "#1E5F9F", lang: str = "id") -> bytes:
    lg = lang if lang in _L else "id"
    t = _L[lg]
    ac = _accent(accent)
    doc = Document()

    # --- Letterhead ---
    head = doc.add_paragraph()
    r = head.add_run(data.org_name)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ac

    title = doc.add_paragraph()
    tr = title.add_run(data.report_title)
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = ac

    sub = doc.add_paragraph()
    sub.add_run(f"{data.engagement_name}").bold = True
    doc.add_paragraph(f"{t['client']}: {data.client_name}")
    doc.add_paragraph(f"{t['prepared_by']}: {data.org_name}")
    doc.add_paragraph(f"{t['generated']}: {data.generated_at}")
    if data.posture:
        posture_label = _POSTURE[lg].get(data.posture, data.posture)
        doc.add_paragraph(f"{t['posture']}: {posture_label}")

    # --- Ringkasan eksekutif ---
    if data.summary_overview or data.summary_key_risks or data.summary_recommendations:
        _heading(doc, t["exec_summary"], ac, level=1)
        if data.summary_stale_note:
            # Ditulis sebelum isi ringkasan agar pembaca tahu duluan bahwa
            # angka di paragraf berikutnya mungkin tak lagi sesuai tabel temuan.
            warn = doc.add_paragraph()
            run = warn.add_run(f"⚠ {data.summary_stale_note}")
            run.italic = True
        if data.summary_overview:
            _labelled(doc, t["overview"], data.summary_overview)
        if data.summary_key_risks:
            _labelled(doc, t["key_risks"], data.summary_key_risks)
        if data.summary_recommendations:
            _labelled(doc, t["recommendations"], data.summary_recommendations)

    # --- Temuan ---
    _heading(doc, f"{t['findings']} ({data.total})", ac, level=1)
    if data.severity_counts:
        counts = ", ".join(f"{k}={v}" for k, v in data.severity_counts.items())
        doc.add_paragraph(f"{t['total']}: {data.total} · {t['by_severity']}: {counts}")

    if not data.findings:
        doc.add_paragraph(t["empty"])
    else:
        for i, f in enumerate(data.findings, start=1):
            _heading(doc, f"{i}. {f.title}", ac, level=2)
            meta = [
                f"{t['severity']}: {f.severity}",
                f"{t['priority']}: P{f.priority}" if f.priority else "",
                f"CVSS: {f.cvss_score}" if f.cvss_score is not None else "",
                f"CWE: {f.cwe}" if f.cwe else "",
                f"OWASP: {f.owasp}" if f.owasp else "",
                f"CVE: {', '.join(f.cve)}" if f.cve else "",
                f"{t['status']}: {f.status}" + (f" ({t['edited']})" if f.edited else ""),
            ]
            doc.add_paragraph(" · ".join(m for m in meta if m))
            _labelled(doc, t["description"], f.description or "—")
            _labelled(doc, t["impact"], f.impact or "—")
            _labelled(doc, t["recommendation"], f.recommendation or "—")

    # --- Footer kerahasiaan ---
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = t["confidential"]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _heading(doc: DocumentT, text: str, color: RGBColor, *, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = color


def _labelled(doc: DocumentT, label: str, body: str) -> None:
    p = doc.add_paragraph()
    lr = p.add_run(f"{label}: ")
    lr.bold = True
    p.add_run(body)
